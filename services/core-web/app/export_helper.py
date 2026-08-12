import io
import os
from minio import Minio
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table as RLTable, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Konfigurasi Koneksi MinIO
MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "127.0.0.1:9000")
MINIO_ACCESS_KEY = os.getenv("MINIO_ROOT_USER", "admin")
MINIO_SECRET_KEY = os.getenv("MINIO_ROOT_PASSWORD", "admin12345")
BUCKET_NAME = "balmon-measurements"
def get_minio_client():
    return Minio(MINIO_ENDPOINT, access_key=MINIO_ACCESS_KEY, secret_key=MINIO_SECRET_KEY, secure=False)
def download_image_buffer(minio_path):
    """Mengunduh gambar dari MinIO dan mengembalikannya sebagai BytesIO buffer"""
    if not minio_path:
        return None
    try:
        client = get_minio_client()
        obj_name = minio_path.replace(f"{BUCKET_NAME}/", "")
        response = client.get_object(BUCKET_NAME, obj_name)
        return io.BytesIO(response.read())
    except Exception as e:
        print(f"Gagal mengunduh gambar {minio_path} dari MinIO: {e}")
        return None
def generate_word_report(sessions_data):
    """
    Membuat dokumen Word (.docx) berformat resmi Balmon SFR untuk beberapa stasiun radio.
    `sessions_data` berisi list tuple (MeasurementSession, StasiunRadio, HasilPengukuran)
    """
    doc = Document()
    # Mengatur margin halaman 0.8 inci
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    for idx, (sesi, stasiun, hasil) in enumerate(sessions_data):
        # 1. Judul Dokumen Utama
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(f"HASIL PENGUKURAN KARAKTERISTIK STASIUN RADIO SIARAN FM\nWILAYAH LAYANAN {stasiun.kab_kota.upper()}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        doc.add_paragraph()  # Spacer
        # 2. Tabel Ringkasan Parameter Stasiun Radio
        table = doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        items = [
            ("Nama Radio", f"{stasiun.nama_penyelenggara}, PT. ({stasiun.nama_stasiun})"),
            ("Alamat", f"{stasiun.kab_kota}"),
            ("Frekuensi (MHz)", f"{stasiun.frekuensi_mhz}"),
            ("Level (dBm)", f"{hasil.level_dbm if hasil else '-'}"),
            ("Bandwidth (KHz)", f"{hasil.band_width_khz if hasil else '-'}"),
            ("Deviasi (KHz)", f"{hasil.deviasi_khz if hasil else '-'}")
        ]
        for i, (label, val) in enumerate(items):
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = val
            
            p0 = row_cells[0].paragraphs[0]
            p0.runs[0].font.bold = True
            p0.runs[0].font.size = Pt(10)
            p0.runs[0].font.name = 'Arial'
            p1 = row_cells[1].paragraphs[0]
            p1.runs[0].font.size = Pt(10)
            p1.runs[0].font.name = 'Arial'
        doc.add_paragraph()  # Spacer
        # 3. Sisipkan Foto Pengukuran Spectrum Analyzer
        image_captions = [
            (sesi.obw_img_path, "Capture Center Frekuensi & Band Width"),
            (sesi.deviasi_img_path, "Capture Deviasi Frekuensi"),
            (sesi.harmonisa_img_path, "Capture Harmonisa 1, 2, 3")
        ]
        for img_path, caption in image_captions:
            img_buf = download_image_buffer(img_path)
            if img_buf:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(2)
                p_img.add_run().add_picture(img_buf, width=Inches(5.6))
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_cap = p_cap.add_run(caption)
                r_cap.font.size = Pt(9)
                r_cap.font.italic = True
                r_cap.font.name = 'Arial'
                p_cap.paragraph_format.space_after = Pt(10)
        # 4. Catatan Analisis Hukum & Audit AI (LLM) di Akhir Setiap Stasiun Radio
        if hasil and hasil.catatan_llm:
            p_llm_head = doc.add_paragraph()
            p_llm_head.paragraph_format.space_before = Pt(12)
            r_llm = p_llm_head.add_run("Catatan Analisis Audit Teknis & Evaluasi Hukum AI (Local LLM):")
            r_llm.bold = True
            r_llm.font.size = Pt(10.5)
            r_llm.font.name = 'Arial'
            r_llm.font.color.rgb = RGBColor(0, 102, 51)  # Hijau Balmon
            p_llm_body = doc.add_paragraph()
            p_llm_body.paragraph_format.space_before = Pt(3)
            r_body = p_llm_body.add_run(hasil.catatan_llm)
            r_body.font.size = Pt(9)
            r_body.font.name = 'Arial'
        # Halaman Baru untuk Stasiun Berikutnya
        if idx < len(sessions_data) - 1:
            doc.add_page_break()
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
def generate_pdf_report(sessions_data):
    """
    Membuat dokumen PDF berformat resmi Balmon SFR untuk beberapa stasiun radio.
    """
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    story = []
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        alignment=1,
        spaceAfter=12
    )
    cell_bold = ParagraphStyle('CellBold', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8.5, leading=10)
    cell_norm = ParagraphStyle('CellNorm', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=10)
    caption_style = ParagraphStyle('CaptionStyle', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=8.5, leading=10, alignment=1, spaceAfter=8)
    llm_head_style = ParagraphStyle('LLMHead', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=9.5, leading=12, textColor=colors.HexColor('#006633'), spaceBefore=8, spaceAfter=3)
    llm_body_style = ParagraphStyle('LLMBody', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=10, spaceAfter=8)
    for idx, (sesi, stasiun, hasil) in enumerate(sessions_data):
        # 1. Judul
        title_text = f"HASIL PENGUKURAN KARAKTERISTIK STASIUN RADIO SIARAN FM<br/>WILAYAH LAYANAN {stasiun.kab_kota.upper()}"
        story.append(Paragraph(title_text, title_style))
        # 2. Tabel Parameter Stasiun
        data_table = [
            [Paragraph("Nama Radio", cell_bold), Paragraph(f"{stasiun.nama_penyelenggara}, PT. ({stasiun.nama_stasiun})", cell_norm)],
            [Paragraph("Alamat", cell_bold), Paragraph(f"{stasiun.kab_kota}", cell_norm)],
            [Paragraph("Frekuensi (MHz)", cell_bold), Paragraph(f"{stasiun.frekuensi_mhz}", cell_norm)],
            [Paragraph("Level (dBm)", cell_bold), Paragraph(f"{hasil.level_dbm if hasil else '-'}", cell_norm)],
            [Paragraph("Bandwidth (KHz)", cell_bold), Paragraph(f"{hasil.band_width_khz if hasil else '-'}", cell_norm)],
            [Paragraph("Deviasi (KHz)", cell_bold), Paragraph(f"{hasil.deviasi_khz if hasil else '-'}", cell_norm)]
        ]
        t = RLTable(data_table, colWidths=[140, 380])
        t.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ('TOPPADDING', (0,0), (-1,-1), 3),
        ]))
        story.append(t)
        story.append(Spacer(1, 8))
        # 3. Sisipkan Foto Pengukuran Spectrum Analyzer
        image_captions = [
            (sesi.obw_img_path, "Capture Center Frekuensi & Band Width"),
            (sesi.deviasi_img_path, "Capture Deviasi Frekuensi"),
            (sesi.harmonisa_img_path, "Capture Harmonisa 1, 2, 3")
        ]
        for img_path, caption in image_captions:
            img_buf = download_image_buffer(img_path)
            if img_buf:
                try:
                    rl_img = RLImage(img_buf, width=460, height=230)
                    story.append(rl_img)
                    story.append(Spacer(1, 2))
                    story.append(Paragraph(caption, caption_style))
                except Exception as e:
                    print(f"Error render PDF image: {e}")
        # 4. Catatan Analisis LLM di Akhir Stasiun
        if hasil and hasil.catatan_llm:
            story.append(Paragraph("Catatan Analisis Audit Teknis & Evaluasi Hukum AI (Local LLM):", llm_head_style))
            formatted_llm_text = hasil.catatan_llm.replace("\n", "<br/>")
            story.append(Paragraph(formatted_llm_text, llm_body_style))
        if idx < len(sessions_data) - 1:
            story.append(PageBreak())
    doc.build(story)
    output.seek(0)
    return output